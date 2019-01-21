# coding:utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt
import docx.enum.text as alg

doc = Document()
run = doc.add_paragraph().add_run()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.bold = True

baslik = doc.add_paragraph("UZLAŞTIRMA RAPORU")
baslik.alignment = alg.WD_ALIGN_PARAGRAPH.CENTER
baslik.bold = False
baslik.add_run()

doc.save("deneme.docx")
# -*- coding:utf-8 -*-

from vtbgln import VbagKur
import datetime
import math, time, os
from shutil import copy2
import xlwings as xw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from PyQt5 import QtWidgets
import sys

db = VbagKur()

def kolon_tara(sutun, tablo):
    sor = db.kolon_oku(sutun, tablo)
    return sor

def dosya_durumu_tara(dosya):
    arg = db.hepsini_oku("durum", "dosyalar", "uzno", dosya)
    if arg[0][0] == 0:
        return ("Dosya Oluşturuldu.")
    elif arg[0][0] == 1:
        return ("Uzlaşma aşamasında")
    elif arg[0][0] == 4:
        return ("Uzlaşma sonlandı")

def tekli_demet_coz(demet):
    dizi = []
    for i in range(len(demet)):
        dizi.append(demet[i][0])
    return dizi

def tek_satirlik_demet_coz(demet):
    dizi = []
    for i in range(len(demet[0])):
        dizi.append(demet[0][i])
    return dizi

def dosya_cek(dosya):
    sor = db.hepsini_oku("*", "dosyalar", "uzno", dosya)
    return sor

def dosya_uzatma_kontrol(dosya):
    sor = db.komut("select * from uzatma where dosya = '{}'".format(dosya))
    if len(sor) == 0:
        return False
    else:
        return True

def kmt(dosya):
    son = db.komut(dosya)
    return son

def silinecek_veri_bul(arg, tablo):
    sonuc = tek_satirlik_demet_coz(db.komut(arg))
    db.satir_sil("DELETE FROM {} WHERE id = '{}'".format(tablo, sonuc[0]))
    return True

def silinecek_gider_bul(arg):
    sonuc = tek_satirlik_demet_coz(db.komut(arg))
    db.satir_sil("DELETE FROM giderler WHERE id = '{}'".format(sonuc[0]))
    return True

def silinecek_ek_bul(arg):
    sonuc = tek_satirlik_demet_coz(db.komut(arg))
    db.satir_sil("DELETE FROM ek WHERE id = '{}'".format(sonuc[0]))
    return True

def tarih_duzenle(arg=None):
    if arg == None:
        return False
    else:
        t = arg.replace("-","/")
        formatstring = '%Y/%m/%d'
        dosya_bitis = datetime.datetime.strptime(t, formatstring)
        bugun = datetime.datetime.now()
        fark = dosya_bitis - bugun
        return str(fark)

def uzatma_tarihi_ekle(arg, dosya):
    sql = """UPDATE dosyalar SET 
        uzatmatar='{}' WHERE uzno = '{}'""".format(arg, dosya)
    sor = db.komut(sql)

def uzlastirma_suresi_bul(baslama, bitis):
    bt = baslama.replace(".", "/")
    formatstr = '%d/%m/%Y'
    bt2 = datetime.datetime.strptime(bt, formatstr)

    bbt = bitis.replace(".", "/")
    formatstr = '%d/%m/%Y'
    bbt2 = datetime.datetime.strptime(bbt, formatstr)

    fark_s = bt2 - bbt2
    farstr = str(fark_s).split()
    return abs(int(farstr[0]))

def olay_ozeti_cek(arg):
    return db.komut("select ozet from olaylar where uzno = '{}'".format(arg))

def gorusme_cek(arg):
    return db.komut("select gorusme from uzgor where dosya = '{}'".format(arg))

def edim_cek(arg):
    return db.komut("select edi from edim where dosya = '{}'".format(arg))

def uzbas_cek(arg):
    return db.komut("select sebeb from uzbas where dosya = '{}'".format(arg))

def sonuc_cek(arg):
    return db.komut("select metin from sonuc where dosya = '{}'".format(arg))


def klasor_dogrula(yol, mesaj, baslik):
    if os.path.isdir(yol) == True:
        return True
    else:
        inf = QtWidgets.QMessageBox()
        inf.setIcon(QtWidgets.QMessageBox.Information)
        inf.setWindowTitle(baslik)
        inf.setText(mesaj)
        inf.setStandardButtons(QtWidgets.QMessageBox.Ok)
        inf.exec_()
        return False

def davet_yaz(durum, sahis, tc, veksicil, no, ttarihi, uz, uzsicil, uzte):
    ana_dizin = os.getcwd()
    ana_dosya_yolu = "\\dosyalar\\uzfile\\Davet Mektubu.xlsx"
    sor = db.komut("select kayit_yeri from ayarlar")
    if klasor_dogrula(sor[0][0],
                   "Dosyaları oluşturabilmek için dosya kayıt yeri belirlemeniz gerek."
                   "Ayarlar'dan geçerli bir dizin seçiniz",
                   "Dosya Kayıt Yeri Hatası") == False:
        return
    else:
        pass

    sorno = no.replace("/", "-")
    hedef_dizin = sor[0][0] + "\\" + sorno
    try:
        os.mkdir(hedef_dizin)
    except FileExistsError:
        pass
    kaynak = ana_dizin + ana_dosya_yolu
    hedef = hedef_dizin +"\\" +sahis + " Davet Mektubu.xlsx"
    copy2(kaynak, hedef)

    wb = xw.Book(hedef)
    sht = wb.sheets['Sayfa1']
    sht.range('AT3').value = durum
    sht.range('AT4').value = sahis
    sht.range('AT5').value = tc
    sht.range('AT6').value = no
    sht.range('AT7').value = ttarihi
    sht.range('AT8').value = uz
    sht.range('AT9').value = uzsicil
    sht.range('AT10').value = uzte
    sht.range('AT11').value = veksicil
    return True

def teklif_yaz(sorno, ad, uz, uzsicil):
    sor = db.komut("select * from taraflar where dosya='{}' and ad='{}'".format(sorno, ad))
    esas_sor = db.komut("select mahesno from dosyalar where uzno='{}'".format(sorno))

    ttarihi = sor[0][2]
    sifat = sor[0][3]
    tc = sor[0][4]
    badı = sor[0][5]
    aadı = sor[0][6]
    dyeri = sor[0][7]
    dtar = sor[0][8]
    adres = sor[0][12]

    ana_dizin = os.getcwd()
    ana_dosya_yolu = "\\dosyalar\\uzfile\\Teklif Formu.xlsx"
    yersor = db.komut("select kayit_yeri from ayarlar")
    if klasor_dogrula(yersor[0][0],
                   "Dosyaları oluşturabilmek için dosya kayıt yeri belirlemeniz gerek."
                   "Ayarlar'dan geçerli bir dizin seçiniz",
                   "Dosya Kayıt Yeri Hatası") == False:
        return
    else:
        pass
    no = sorno.replace("/", "-")
    hedef_dizin = yersor[0][0] + "\\" + no
    try:
        os.mkdir(hedef_dizin)
    except FileExistsError:
        pass
    kaynak = ana_dizin + ana_dosya_yolu
    hedef = hedef_dizin +"\\" +ad + " Teklif Formu.xlsx"
    copy2(kaynak, hedef)

    kanun = db.komut("select icerik from sablonlar where id = '1'")
    wb = xw.Book(hedef)
    sht = wb.sheets['Sayfa1']
    sht.range('AU1').value = kanun[0][0]
    sht.range('AU2').value = ttarihi
    sht.range('AU3').value = uz
    sht.range('AU4').value = uzsicil
    sht.range('AU6').value = tc
    sht.range('AU7').value = ad
    sht.range('AU8').value = badı
    sht.range('AU9').value = aadı
    sht.range('AU10').value = dyeri
    sht.range('AU11').value = dtar
    sht.range('AU12').value = adres
    sht.range('AU19').value = sorno
    sht.range('AU20').value = esas_sor[0][0]
    if sifat == 1:
        sht.range('AU13').value = "( X )"
    elif sifat == 2:
        sht.range('AU14').value = "( X )"
    elif sifat == 3:
        sht.range('AU15').value = "( X )"
    elif sifat == 4:
        sht.range('AU16').value = "( X )"
    elif sifat == 5:
        sht.range('AU17').value = "( X )"
    elif sifat == 6:
        sht.range('AU18').value = "( X )"
    else:
        print("Sicil durumu veya nitelik durumunda hata var.")
    return True

def tebligat_yaz(sorno, ad):
    sor = db.komut("select * from taraflar where dosya='{}' and ad='{}'".format(sorno, ad))
    ttarihi = sor[0][2]
    sifat = sor[0][3]
    tc = sor[0][4]
    badı = sor[0][5]
    aadı = sor[0][6]
    dyeri = sor[0][7]
    dtar = sor[0][8]
    cinsiyet = sor[0][9]
    adres = sor[0][12]

    ana_dizin = os.getcwd()
    ana_dosya_yolu = "\\dosyalar\\uzfile\\Tebligat.xlsx"
    yersor = db.komut("select kayit_yeri from ayarlar")
    if klasor_dogrula(yersor[0][0],
                   "Dosyaları oluşturabilmek için dosya kayıt yeri belirlemeniz gerek."
                   "Ayarlar'dan geçerli bir dizin seçiniz",
                   "Dosya Kayıt Yeri Hatası") == False:
        return
    else:
        pass
    no = sorno.replace("/", "-")
    hedef_dizin = yersor[0][0] + "\\" + no
    try:
        os.mkdir(hedef_dizin)
    except FileExistsError:
        pass
    kaynak = ana_dizin + ana_dosya_yolu
    hedef = hedef_dizin +"\\" +ad + " Tebligat.xlsx"
    copy2(kaynak, hedef)

    wb = xw.Book(hedef)
    sht = wb.sheets['Sayfa1']
    sht.range('Q1').value = sorno
    sht.range('Q2').value = ad
    sht.range('Q3').value = badı
    sht.range('Q4').value = aadı
    sht.range('Q5').value = cinsiyet
    sht.range('Q6').value = dtar
    sht.range('Q7').value = tc
    sht.range('Q8').value = adres
    sht.range('Q9').value = ttarihi
    sorno_sor = db.komut("select sorno, mahesno from dosyalar where uzno='{}'".format(sorno))
    if sorno_sor[0][0] == "":
        sht.range('Q10').value = "Mahkeme"
    else:
        sht.range('Q10').value = "Savcılık"
    return True

def rapor_yaz(sorno, uz, uz_sicil, rapor_yeri):
    yersor = db.komut("select kayit_yeri from ayarlar")
    if klasor_dogrula(yersor[0][0],
                   "Dosyaları oluşturabilmek için dosya kayıt yeri belirlemeniz gerek."
                   "Ayarlar'dan geçerli bir dizin seçiniz",
                   "Dosya Kayıt Yeri Hatası") == False:
        return
    else:
        pass
    no = sorno.replace("/", "-")
    hedef_dizini = yersor[0][0] + "\\" + no
    try:
        os.mkdir(hedef_dizini)
    except FileExistsError:
        pass

    hedef_dizin = yersor[0][0] + "/" + no + "/" + "rapor.docx"

    dosya = db.komut("select * from dosyalar where uzno = '{}'".format(sorno))
    uzlastirma_no = dosya[0][1]
    sorusturma_no = dosya[0][2]
    uz_bilgileri = db.komut("select adres from uzlasmaci where isim = '{}'".format(uz))
    uz_adres = uz_bilgileri[0][0]

    doc = Document("dosyalar/default.docx")

    run = doc.add_paragraph().add_run()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.bold = True

    section = doc.sections[0]
    footer = section.footer
    yl = footer.paragraphs[0]
    yli = yl.add_run("\nT.C. İstanbul Cumhuriyet Başsavcılığı Uzlaştırma Bürosu Uzl.No: ")
    yli.italic = True
    yli.font.size = Pt(9)
    yli.font.bold = False
    yli = yl.add_run(uzlastirma_no)
    yli.font.color.rgb = RGBColor(255, 0, 0)
    yli.italic = True
    yli.font.size = Pt(9)
    yli.font.bold = False

    ylll = footer.paragraphs[0]
    yli = ylll.add_run("\nT.C. İstanbul Cumhuriyet Başsavcılığı Soruşturma No: ")
    yli.italic = True
    yli.font.size = Pt(9)
    yli.font.bold = False
    yli = yl.add_run(sorusturma_no)
    yli.font.color.rgb = RGBColor(255, 0, 0)
    yli.italic = True
    yli.font.size = Pt(9)
    yli.font.bold = False
    yli.left_indent = Inches(0.25)

    baslik = doc.add_paragraph("UZLAŞTIRMA RAPORU")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.add_run()

    uzno = doc.add_paragraph("Uzlaştırma No\t\t\t\t:{}".format(dosya[0][1]))
    uzno.add_run()

    sno = doc.add_paragraph("Cumhuriyet Başsavcılığı Soruşturma No\t:{}".format(dosya[0][2]))
    sno.add_run()

    esno = doc.add_paragraph("Mahkeme Esas No\t\t\t\t:{}".format(dosya[0][3]))
    esno.add_run()

    suc = doc.add_paragraph("Uzlaştırma Konusu Suç/Suçlar \t\t:{}".format(dosya[0][4]))
    suc.add_run()

    doc_uz = doc.add_paragraph("Uzlaştırmacının")
    doc_uz.paragraph_format.space_before = Pt(12)
    doc_uz.add_run()

    adi = doc.add_paragraph()
    adi.paragraph_format.left_indent = Inches(0.5)
    i_adi = adi.add_run("Adı Soyadı \t\t\t\t:{}".format(uz))
    i_adi.bold = False

    sici = doc.add_paragraph()
    sici.paragraph_format.left_indent = Inches(0.5)
    i_sici = sici.add_run("Uzlaştırmacının Sicili \t\t:{}".format(uz_sicil))
    i_sici.bold = False

    dos_uzad = doc.add_paragraph()
    dos_uzad.paragraph_format.left_indent = Inches(0.5)
    i_dos_uzad = dos_uzad.add_run("İletişim Adresi \t\t\t:{}".format(uz_bilgileri[0][0]))
    i_dos_uzad.bold = False

    gortar = db.komut("select tektar from dosyalar where uzno = '{}'".format(sorno))
    doc_gorev_tarihi = doc.add_paragraph("Görevlendirme Tarihi\t\t\t:{}".format(gortar[0][0]))
    doc_gorev_tarihi.add_run()

    doc_evrak_tarihi = doc.add_paragraph("Dosya içindeki belgelerin birer Örneğinin")
    doc_evrak_tarihi.add_run()
    doc_evrak_tarihi = doc.add_paragraph(" verildiği Uzl. Süresinin başladığı tarih\t:{}".format(gortar[0][0]))
    doc_evrak_tarihi.add_run()

    ilkteklif = db.komut("""select ttarihi from temsilciler where dosya='{0}' union select ttarihi from taraflar
                         where dosya='{0}' order by ttarihi asc """.format(sorno))
    doc_uzlasmateklif = doc.add_paragraph("Uzlaşma Teklif Tarihi\t\t\t:{}".format(ilkteklif[0][0]))
    doc_uzlasmateklif.add_run()

    uzatma = db.komut("select tarih from uzatma where dosya = '{}'".format(sorno))
    uzsure = db.komut("select uzatma_suresi from ayarlar")
    if len(uzatma) == 0:
        doc_eksure = doc.add_paragraph("Ek süre verilme tarihi ve süresi\t\t:{}".format(" "))
        doc_eksure.paragraph_format.space_after = Pt(12)
        doc_eksure.add_run()
    else:
        doc_eksure = doc.add_paragraph("Ek süre verilme tarihi ve süresi\t\t:{} / {} gün".format(uzatma[0][0], uzsure[0][0]))
        doc_eksure.paragraph_format.space_after = Pt(12)
        doc_eksure.add_run()


    saniklar = db.komut("select * from taraflar where dosya='{}'and sifat='5'".format(sorno))
    sanik_temsilcileri = db.komut("select * from taraflar where dosya='{}'and sifat='6'".format(sorno))
    magdurlar = db.komut("select * from taraflar where dosya='{}'and sifat='1'".format(sorno))
    magdur_temsilcileri = db.komut("select * from taraflar where dosya='{}'and sifat='2'".format(sorno))
    zarar_gorenler = db.komut("select * from taraflar where dosya='{}'and sifat='3'".format(sorno))
    zarar_goren_temsilcileri = db.komut("select * from taraflar where dosya='{}'and sifat='4'".format(sorno))
    mustekiler = db.komut("select * from taraflar where dosya='{}'and sifat='7'".format(sorno))
    musteki_temsilcileri = db.komut("select * from taraflar where dosya='{}'and sifat='8'".format(sorno))

    #rapor imza kısmında kullanılacakkisi ve sifatları tutacak boş liste
    kisi_listesi = []
    #------------------------------------------------ŞÜPHELİ------------
    if len(saniklar) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Şüphelinin /Sanığın / Kanuni Temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(saniklar)):
            tup_liste = (saniklar[i][1], "5")
            kisi_listesi.append(tup_liste)

            s_ad = doc.add_paragraph()
            s_ad.paragraph_format.left_indent = Inches(0.5)
            i_s_ad = s_ad.add_run("Ad Soyad \t\t:{}".format(saniklar[i][1]))
            i_s_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(saniklar[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(saniklar[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'"
                                 .format(sorno, saniklar[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    tup_liste_vekil = (vekil_ara[im][2], "Müdafii")
                    kisi_listesi.append(tup_liste_vekil)

                    doc_s1 = doc.add_paragraph("Müdafiin")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()

                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    vv_adr = doc.add_paragraph()
                    vv_adr.paragraph_format.left_indent = Inches(0.5)
                    i_vv_adr = vv_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_vv_adr.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    v_adr.paragraph_format.space_after = Pt(12)
                    i_v_adr = v_adr.add_run("Kayıtlı olduğu Baro ve Sicil No \t:{} / {}"
                                            .format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_v_adr.bold = False
        # ------------------------------------------------ŞÜPHELİ TEMSİLCİSİ------------
    if len(sanik_temsilcileri) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Şüphelinin /Sanığın / Kanuni Temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(sanik_temsilcileri)):
            tup_sanik = (sanik_temsilcileri[i][1], "6")
            kisi_listesi.append(tup_sanik)

            s_ad = doc.add_paragraph()
            s_ad.paragraph_format.left_indent = Inches(0.5)
            i_s_ad = s_ad.add_run("Ad Soyad \t\t:{}".format(sanik_temsilcileri[i][1]))
            i_s_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(sanik_temsilcileri[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(sanik_temsilcileri[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut(
                "select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, sanik_temsilcileri[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    tup_san_vekil = (vekil_ara[im][2], "Müdafii")
                    kisi_listesi.append(tup_san_vekil)

                    doc_s1 = doc.add_paragraph("Müdafiin")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()

                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    vv_adr = doc.add_paragraph()
                    vv_adr.paragraph_format.left_indent = Inches(0.5)
                    i_vv_adr = vv_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_vv_adr.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    v_adr.paragraph_format.space_after = Pt(12)
                    i_v_adr = v_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_v_adr.bold = False
    #------------------------------------------------MAĞDUR
    if len(magdurlar) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Mağdur / Katılan / Suçtan zarar görenin /Kanunî temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(magdurlar)):
            s_ad = doc.add_paragraph()
            s_ad.paragraph_format.left_indent = Inches(0.5)
            i_s_ad = s_ad.add_run("Ad Soyad \t\t:{}".format(magdurlar[i][1]))
            i_s_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(magdurlar[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(magdurlar[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, magdurlar[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    doc_s1 = doc.add_paragraph("Vekilin")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()

                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    vs_adr = doc.add_paragraph()
                    vs_adr.paragraph_format.left_indent = Inches(0.5)
                    i_vs_adr = vs_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_vs_adr.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    v_adr.paragraph_format.space_after = Pt(12)
                    i_v_adr = v_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_v_adr.bold = False
    #----------------------------------------------------------------MAĞDUR tEMSİLCİSİ
    if len(magdur_temsilcileri) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Mağdur / Katılan / Suçtan zarar görenin /Kanunî temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(magdur_temsilcileri)):
            t_ad = doc.add_paragraph()
            t_ad.paragraph_format.left_indent = Inches(0.5)
            i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(magdur_temsilcileri[i][1]))
            i_t_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(magdur_temsilcileri[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(magdur_temsilcileri[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, magdur_temsilcileri[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    doc_s1 = doc.add_paragraph("Vekil")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()
                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    i_v_adr = v_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_v_adr.bold = False

                    vb_adr = doc.add_paragraph()
                    vb_adr.paragraph_format.left_indent = Inches(0.5)
                    vb_adr.paragraph_format.space_after = Pt(12)
                    i_vb_adr = vb_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_vb_adr.bold = False
    #---------------------------------------------------------Zarar Görenler
    if len(zarar_gorenler) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Suçtan Zarar Gören / Kanuni Temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(zarar_gorenler)):
            t_ad = doc.add_paragraph()
            t_ad.paragraph_format.left_indent = Inches(0.5)
            i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(zarar_gorenler[i][1]))
            i_t_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(zarar_gorenler[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(zarar_gorenler[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, zarar_gorenler[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    doc_s1 = doc.add_paragraph("Vekil")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()
                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    i_v_adr = v_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_v_adr.bold = False

                    vb_adr = doc.add_paragraph()
                    vb_adr.paragraph_format.left_indent = Inches(0.5)
                    vb_adr.paragraph_format.space_after = Pt(12)
                    i_vb_adr = vb_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_vb_adr.bold = False
    #--------------------------------------Zarar Görenlerin Temsilcileri
    if len(zarar_goren_temsilcileri) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Suçtan Zarar Gören / Kanuni Temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(zarar_goren_temsilcileri)):
            t_ad = doc.add_paragraph()
            t_ad.paragraph_format.left_indent = Inches(0.5)
            i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(zarar_goren_temsilcileri[i][1]))
            i_t_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(zarar_goren_temsilcileri[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(zarar_goren_temsilcileri[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, zarar_goren_temsilcileri[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    doc_s1 = doc.add_paragraph("Vekil")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()
                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    i_v_adr = v_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_v_adr.bold = False

                    vb_adr = doc.add_paragraph()
                    vb_adr.paragraph_format.left_indent = Inches(0.5)
                    vb_adr.paragraph_format.space_after = Pt(12)
                    i_vb_adr = vb_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_vb_adr.bold = False

    #-----------------------------------Müşteki Şüpheli
    if len(mustekiler) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Müşteki Şüphelinin / Kanuni Temsilcisinin")
        doc_s1.paragraph_format.space_before = Pt(12)
        doc_s1.add_run()
        for i in range(len(mustekiler)):
            t_ad = doc.add_paragraph()
            t_ad.paragraph_format.left_indent = Inches(0.5)
            i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(mustekiler[i][1]))
            i_t_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(mustekiler[i][4]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(mustekiler[i][12]))
            i_s_adr.bold = False

            vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno, mustekiler[i][1]))
            if len(vekil_ara) == 0:
                pass
            else:
                for im in range(len(vekil_ara)):
                    doc_s1 = doc.add_paragraph("Müdafinin")
                    doc_s1.paragraph_format.space_before = Pt(12)
                    doc_s1.add_run()
                    v_ad = doc.add_paragraph()
                    v_ad.paragraph_format.left_indent = Inches(0.5)
                    i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                    i_v_ad.bold = False

                    v_tc = doc.add_paragraph()
                    v_tc.paragraph_format.left_indent = Inches(0.5)
                    i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                    i_v_tc.bold = False

                    v_adr = doc.add_paragraph()
                    v_adr.paragraph_format.left_indent = Inches(0.5)
                    i_v_adr = v_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                    i_v_adr.bold = False

                    vb_adr = doc.add_paragraph()
                    vb_adr.paragraph_format.left_indent = Inches(0.5)
                    vb_adr.paragraph_format.space_after = Pt(12)
                    i_vb_adr = vb_adr.add_run(
                        "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                    i_vb_adr.bold = False
        #--------------Müşteki temsilcileri
        if len(musteki_temsilcileri) == 0:
            pass
        else:
            print(musteki_temsilcileri)
            doc_s1 = doc.add_paragraph("Müşteki Şüphelinin / Kanuni Temsilcisinin")
            doc_s1.paragraph_format.space_before = Pt(12)
            doc_s1.add_run()
            for i in range(len(musteki_temsilcileri)):
                t_ad = doc.add_paragraph()
                t_ad.paragraph_format.left_indent = Inches(0.5)
                i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(musteki_temsilcileri[i][1]))
                i_t_ad.bold = False

                s_tc = doc.add_paragraph()
                s_tc.paragraph_format.left_indent = Inches(0.5)
                i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(musteki_temsilcileri[i][4]))
                i_s_tc.bold = False

                s_adr = doc.add_paragraph()
                s_adr.paragraph_format.left_indent = Inches(0.5)
                s_adr.paragraph_format.space_after = Pt(12)
                i_s_adr = s_adr.add_run("Adres \t\t:{}".format(musteki_temsilcileri[i][12]))
                i_s_adr.bold = False

                vekil_ara = db.komut("select * from temsilciler where dosya='{}'and kisi='{}'".format(sorno,
                                                                                                      musteki_temsilcileri[
                                                                                                          i][1]))
                if len(vekil_ara) == 0:
                    pass
                else:
                    for im in range(len(vekil_ara)):
                        doc_s1 = doc.add_paragraph("Müdafinin")
                        doc_s1.paragraph_format.space_before = Pt(12)
                        doc_s1.add_run()
                        v_ad = doc.add_paragraph()
                        v_ad.paragraph_format.left_indent = Inches(0.5)
                        i_v_ad = v_ad.add_run("Ad Soyad \t\t:{}".format(vekil_ara[im][2]))
                        i_v_ad.bold = False

                        v_tc = doc.add_paragraph()
                        v_tc.paragraph_format.left_indent = Inches(0.5)
                        i_v_tc = v_tc.add_run("T.C. Kimlik No \t:{}".format(vekil_ara[im][11]))
                        i_v_tc.bold = False

                        v_adr = doc.add_paragraph()
                        v_adr.paragraph_format.left_indent = Inches(0.5)
                        i_v_adr = v_adr.add_run("Adres \t\t:{}".format(vekil_ara[im][5]))
                        i_v_adr.bold = False

                        vb_adr = doc.add_paragraph()
                        vb_adr.paragraph_format.left_indent = Inches(0.5)
                        vb_adr.paragraph_format.space_after = Pt(12)
                        i_vb_adr = vb_adr.add_run(
                            "Kayıtlı olduğu Baro ve Sicil No \t:{} / {}".format(vekil_ara[im][10], vekil_ara[im][3]))
                        i_vb_adr.bold = False

    tercu = db.komut("select * from tercuman where dosya='{}'")
    if len(tercu) == 0:
        pass
    else:
        doc_s1 = doc.add_paragraph("Tercüman")
        doc_s1.add_run()
        for i in range(len(tercu)):
            t_ad = doc.add_paragraph()
            t_ad.paragraph_format.left_indent = Inches(0.5)
            i_t_ad = t_ad.add_run("Ad Soyad \t\t:{}".format(musteki_temsilcileri[i][1]))
            i_t_ad.bold = False

            s_tc = doc.add_paragraph()
            s_tc.paragraph_format.left_indent = Inches(0.5)
            i_s_tc = s_tc.add_run("T.C. Kimlik No \t:{}".format(musteki_temsilcileri[i][2]))
            i_s_tc.bold = False

            s_adr = doc.add_paragraph()
            s_adr.paragraph_format.left_indent = Inches(0.5)
            s_adr.paragraph_format.space_after = Pt(12)
            i_s_adr = s_adr.add_run("Adres \t\t:{}".format(musteki_temsilcileri[i][3]))
            i_s_adr.bold = False

    #Türkiye'de gösterilen adres var mı ?
    sor = db.komut("select adres from taraflar where dosya='{}' and adresniteligi='1'".format(sorno))
    if len(sor) == 0:
        pass
    else:
        for i in range(len(sor)):
            doc_adr1 = doc.add_paragraph("Taraflardan biri yabancı ülkede oturuyorsa \n"
                                         "Türkiye’de göstereceği ikametgâhı\t:{}".format(sor[i][0]))
            doc_adr1.add_run()
    #Yabancı Ülke Adresi
    sor = db.komut("select adres from taraflar where dosya='{}' and adresniteligi='2'".format(sorno))
    if len(sor) == 0:
        pass
    else:
        for i in range(len(sor)):
            doc_adr1 = doc.add_paragraph("Taraflardan biri yabancı ve Türkiye’de\n"
                                         "Göstereceği bir ikametgâhı yok ise \n"
                                         "Ülkesindeki ikametgâhı\t:{}".format(sor[i][0]))
            doc_adr1.add_run()

    bugun_sor = datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    bugun_bol = bugun_sor.split()
    bugun = bugun_bol[0]

    doc_adr1 = doc.add_paragraph("Raporun Düzenlendiği Yer ve Tarih\t:{} / {}".format(rapor_yeri, bugun.replace("/", ".")))
    doc_adr1.add_run()

    fark_sure = uzlastirma_suresi_bul(gortar[0][0], bugun)

    doc_adr1 = doc.add_paragraph("Uzlaştırma Süresi\t\t:{}".format(fark_sure))
    doc_adr1.add_run()

    sor = db.komut("select metin from sonuc where dosya='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph()
    doc_adr1.add_run("Uzlaştırma Sonucu\t\t:").underline = True
    doc_adr1.paragraph_format.space_before = Pt(12)
    if len(sor) == 0:
        pass
    else:
        doc_adr1 = doc.add_paragraph()
        doc_adr1.paragraph_format.space_after = Pt(12)
        i = doc_adr1.add_run(sor[0][0])
        i.bold = False


    run = doc.add_paragraph()
    run.paragraph_format.space_after = Pt(12)
    run.add_run("Açıklamalar\t\t:").underline = True

    sor = db.komut("select ozet from olaylar where uzno='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph()
    doc_adr1.add_run("Olayın Özeti\t\t:").underline = True
    if len(sor) == 0:
        pass
    else:

        doc_adr1 = doc.add_paragraph()
        doc_adr1.paragraph_format.space_after = Pt(12)
        i = doc_adr1.add_run(sor[0][0])
        i.bold = False

    sor = db.komut("select gorusme from uzgor where dosya='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph()
    doc_adr1.add_run("Uzlaşma Görüşmeleri\t\t:").underline = True
    doc_adr1.paragraph_format.space_before = Pt(12)
    if len(sor) == 0:
        pass
    else:

        doc_adr1 = doc.add_paragraph()
        doc_adr1.paragraph_format.space_after = Pt(12)
        i = doc_adr1.add_run(sor[0][0])
        i.bold = False

    sor = db.komut("select edi from edim where dosya='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph("Tarafların üzerinde anlaştıkları edim; ediminin yerine getirilme şekli ve zamanı:")
    doc_adr1.add_run()
    doc_adr1.paragraph_format.space_before = Pt(12)
    if len(sor) == 0:
        pass
    else:

        doc_adr1 = doc.add_paragraph()
        doc_adr1.paragraph_format.space_after = Pt(12)
        i = doc_adr1.add_run(sor[0][0])
        i.bold = False

    sor = db.komut("select sebeb from uzbas where dosya='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph("Uzlaştırmanın başarısızlıkla sonuçlanması hâlinde nedenleri:")
    doc_adr1.add_run()
    doc_adr1.paragraph_format.space_before = Pt(12)
    if len(sor) == 0:
        pass
    else:

        doc_adr1 = doc.add_paragraph()
        doc_adr1.paragraph_format.space_after = Pt(12)
        doc_adr1.bold = False
        i = doc_adr1.add_run(sor[0][0])
        i.bold = False


    gider = db.komut("select aciklama, tutar from giderler where dosya='{}'".format(sorno))
    doc_adr1 = doc.add_paragraph()
    doc_adr1.paragraph_format.space_before = Pt(12)
    i = doc_adr1.add_run("Yapılan Giderler:")
    i.underline = True
    if len(gider) == 0:
        pass
    else:

        for i in range(len(gider)):
            run = doc.add_paragraph()
            run.add_run("{} - {} : {} TL".format(i+1, gider[i][0], gider[i][1]))

    ekler = db.komut("select ad from ek where dosya='{}'".format(sorno))
    doc_a = doc.add_paragraph()
    doc_a.paragraph_format.space_before = Pt(12)
    i = doc_a.add_run("EKLER")
    i.underline = True
    if len(ekler) == 0:
        pass
    else:
        for i in range(len(ekler)):
            run = doc.add_paragraph()
            run.add_run("Ek-{} :".format(i+1)).underline = True
            run.add_run(" {}".format(ekler[i][0]))

        run.paragraph_format.space_after = Pt(12)

    sor = db.komut("""SELECT ad, sifat FROM taraflar WHERE dosya == '{0}' 
        union all SELECT ad, sifat FROM temsilciler WHERE dosya == '{0}'
        union all SELECT ad, sifat FROM tercuman WHERE dosya == '{0}'""".format(sorno))

    table = doc.add_table(rows = len(sor)+2 , cols = 3)
    table.style = "Table Grid"
    table.cell(0,2).text = "İmzalar"
    sayac = 1
    for i in range(len(sor)):
        cell = table.cell(sayac+i, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.rows[sayac + i].height = Pt(50)
        if sor[i][1] == 1:
            cell.text = "Mağdur / Katılan"
        elif sor[i][1] == 2:
            cell.text = "Mağdur / Katılan'ın Temsilcisi"
        elif sor[i][1] == 3:
            cell.text = "Suçtan Zarar Gören"
        elif sor[i][1] == 4:
            cell.text = "Suçtan Zarar Görenin Temsilcisi"
        elif sor[i][1] == 5:
            cell.text = "Şüpheli / Sanık"
        elif sor[i][1] == 6:
            cell.text = "Şüpheli Sanık Temsilcisi"
        elif sor[i][1] == 7:
            cell.text = "Müşteki Şüpheli"
        elif sor[i][1] == 8:
            cell.text = "Müşteki Şüpheli Temsilcisi"
        else:
            cell.text = sor[i][1]
        cell = table.cell(sayac+i, 1)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = sor[i][0]

    cell = table.cell(len(sor) + 1, 0)
    cell.text = "Uzlaştırmacı"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell = table.cell(len(sor) + 1, 1)
    cell.text = uz
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.rows[len(sor)+1].height = Pt(50)

    baslik = doc.add_paragraph("ONAY ŞERHİ")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.paragraph_format.space_before = Pt(100)
    baslik.add_run()

    baslik = doc.add_paragraph("Tarih, Mühür ve İmza")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.add_run()

    baslik = doc.add_paragraph("Cumhuriyet Savcısı / Hakim")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.paragraph_format.space_after = Pt(100)
    baslik.add_run()

    baslik = doc.add_paragraph("ONAYLAMAMA GEREKÇESİ:")
    baslik.add_run()
    baslik.paragraph_format.space_before = Pt(40)
    baslik.paragraph_format.space_after = Pt(140)

    baslik = doc.add_paragraph("Tarih, Mühür ve İmza")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.add_run()

    baslik = doc.add_paragraph("Cumhuriyet Savcısı / Hakim")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    baslik.paragraph_format.space_after = Pt(20)
    baslik.add_run()

    doc.save(hedef_dizin)
    print(kisi_listesi)
    return True